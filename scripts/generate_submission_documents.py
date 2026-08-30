from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "docs" / "submission"
TASK_TEMPLATE = Path(
    r"D:/Chat/V-liaot/xwechat_files/wxid_24p20d8t50xb22_6b4f/msg/attach/"
    r"9e20f478899dc29eb19741386f9343c8/2026-08/Rec/a91e48099f5ac98d/F/0/"
    r"湖南信息职业技术学院 2027届学生毕业设计任.docx"
)
RESULT_TEMPLATE = Path(
    r"D:/Chat/V-liaot/xwechat_files/wxid_24p20d8t50xb22_6b4f/msg/attach/"
    r"9e20f478899dc29eb19741386f9343c8/2026-08/Rec/a91e48099f5ac98d/F/1/"
    r"软件学院毕业设计成果参考规范（软件学院）.docx"
)

TITLE = "基于 FastAPI 与 HeadCore 的本地多模态角色陪伴系统设计与实现"
PROJECT = "HutaoChatCore（HeadCore）"
DATE = "2026年8月23日"
STUDENT = "【待填写：学生姓名】"
STUDENT_ID = "【待填写：学号】"
MAJOR = "人工智能技术应用"
CLASS = "人工智能2401班"
SUPERVISOR = "方丽"


def set_run_font(run, *, name: str = "宋体", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    styles = document.styles
    try:
        normal = styles["Normal"]
    except KeyError:
        normal = styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size in (("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "黑体" if style_name != "Title" else "宋体"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
    try:
        styles["Table Grid"]
    except KeyError:
        styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE)
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = "HutaoChatCore（HeadCore）· 毕业设计材料 · "
        set_run_font(paragraph.runs[0], size=9)
        run = paragraph.add_run()
        set_run_font(run, size=9)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instruction)
        run._r.append(end)


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(30)
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(title)
    set_run_font(run, size=18, bold=True)


def add_center(document: Document, text: str, *, size: int = 12, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_body(document: Document, text: str, *, bold_prefix: str = "") -> None:
    lines = text.splitlines() or [""]
    for line in lines:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.space_after = Pt(6)
        if bold_prefix and line.startswith(bold_prefix):
            set_run_font(paragraph.add_run(bold_prefix), bold=True)
            set_run_font(paragraph.add_run(line[len(bold_prefix) :]))
        else:
            set_run_font(paragraph.add_run(line))


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.first_line_indent = Cm(-0.37)
    paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(paragraph.add_run("• " + text))


def add_code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=9)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7C2D0")
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=10)
    document.add_paragraph()


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    lines = text.splitlines() or [""]
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(paragraph.add_run(line), size=10)


def add_metadata_table(document: Document) -> None:
    rows = [
        ["项目名称", TITLE],
        ["项目简称", PROJECT],
        ["学生姓名", STUDENT],
        ["学号", STUDENT_ID],
        ["专业/班级", f"{MAJOR} / {CLASS}"],
        ["指导教师", SUPERVISOR],
        ["运行边界", "本机/局域网，127.0.0.1:8000，不部署公网"],
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], value)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, size=10, bold=True)
    document.add_paragraph()


def add_task_supplement(document: Document) -> None:
    document.add_page_break()
    add_title(document, f"{PROJECT}项目计划书补充说明")
    add_center(document, "本页用于补充学校任务书合并表格中的详细内容", size=11)

    document.add_heading("一、项目任务拆分", level=1)
    add_table(
        document,
        ["任务域", "实际工作内容", "当前状态", "验收依据"],
        [
            ["核心对话", "HeadCore、ChatService、Provider 路由、流式/非流式回复和长度门禁", "已完成", "pytest tests"],
            ["客户端与认证", "Web Desk、微信小程序、OpenAI 兼容接口、注册登录和 CSRF", "已完成主链路", "接口测试、小程序测试"],
            ["音频与世界工具", "文件音频、条件 ASR/TTS、天气/路线等受控证据工具", "条件可用", "配置开启后本机联调"],
            ["视觉工作台", "白名单标签、时序确认、L1 证据链、L2 场景事实和状态展示", "L1/L2 已完成", "视觉定向测试"],
            ["测试与材料", "全量回归、表面审计、运行手册、计划书和成果书", "已完成当前版本", "929 passed、2 skipped"],
        ],
    )

    document.add_heading("二、实施步骤与时间安排", level=1)
    add_table(
        document,
        ["阶段", "时间", "主要工作", "输出"],
        [
            ["一", "2026.06.30-07.31", "需求、HeadCore/Provider 分层、文本主链路", "可运行 Core"],
            ["二", "2026.08.01-08.31", "认证、会话、条件存储、音频、Web Desk/小程序", "客户端与认证链路"],
            ["三", "2026.09.01-09.30", "世界工具、语音条件能力、运行控制和失败回退", "条件能力适配"],
            ["四", "2026.10.01-10.31", "视觉 L1/L2、工作台展示和隐私边界", "视觉证据链"],
            ["五", "2026.11.01-11.15", "测试、审计、材料和本机/局域网验收", "提交材料"],
        ],
    )

    document.add_heading("三、任务分工与验收责任", level=1)
    add_table(
        document,
        ["责任主体", "负责内容", "边界"],
        [
            ["学生", "代码、测试、视觉边界、审计报告、计划书和成果书", "个人项目，不虚构其他开发成员"],
            ["指导教师方丽", "选题确认、阶段检查、成果审核", "不直接修改代码，不代替技术验收"],
        ],
    )

    document.add_heading("四、交付边界与未完成项", level=1)
    add_table(
        document,
        ["类别", "真实状态"],
        [
            ["已交付", "文本聊天、HeadCore、Web Desk、小程序、认证边界、文件音频、视觉 L1/L2、测试和文档"],
            ["条件可用", "数据库、SMTP、ASR/TTS、世界工具、真实摄像头和固定视觉模型"],
            ["未纳入本次成果", "公网域名/HTTPS、公开运营、远程摄像头上传、通用 VLM 和生产级部署"],
        ],
    )

    document.add_heading("五、主要参考资料", level=1)
    for item in REFERENCES[:8]:
        add_bullet(document, item)


def task_description() -> str:
    return (
        "1.课题说明\n"
        "1.1选题依据：当前角色陪伴类应用通常只负责把用户输入转发给大语言模型，身份、关系、记忆、工具证据和安全边界容易分散在客户端或提示词中，导致回复不稳定、能力状态不透明。\n"
        "1.2选题意义：本课题实现一个本机/局域网运行的 HeadCore 核心，把文本聊天、语音文件识别、语音合成、受控世界证据和本地视觉标签纳入统一的单一人格运行时；通过代码级权限、证据过期、输入校验和回复评估，降低模型幻觉与越权写入风险。\n"
        "2.课题目标：完成 FastAPI Core、Web Desk、小程序客户端、认证与会话、HeadCore 人格/关系/记忆、Provider 路由、文件音频、视觉 L1 标签与规则式 L2 场景事实汇总，并完成测试、运行手册和成果文档。系统只在本机或局域网演示，不进行公网部署。"
    )


def task_requirements() -> str:
    return (
        "1.任务要求\n"
        "（1）完成 FastAPI Core 与 HeadCore 单一人格主链路，支持文本聊天、流式回复、会话和回复评估。\n"
        "（2）完成 Web Desk、原生微信小程序和 OpenAI 兼容文本接口的可运行客户端边界。\n"
        "（3）完成注册、登录、邮箱验证、密码重置、个人中心、CSRF、限流和审计等认证能力。\n"
        "（4）完成文件音频上传、条件 ASR/TTS、受控天气/路线工具和条件数据库适配。\n"
        "（5）完成本地视觉工作台、白名单标签证据链和规则式 L2 场景事实汇总；默认不保存原始帧。\n"
        "（6）补齐自动化测试、运行手册、项目审计报告和毕业设计成果材料。"
    )


def implementation_schedule() -> str:
    return (
        "执行责任：本项目为个人开发，代码、测试、视觉边界和材料整理由学生完成；指导教师方丽负责选题确认、阶段检查和成果审核。\n"
        "第一阶段（2026年6月30日至7月31日）：完成需求梳理、HeadCore/Provider 分层和文本聊天主链路。\n"
        "第二阶段（2026年8月1日至8月31日）：完成认证、会话、双库条件适配、文件音频和 Web Desk/小程序接口。\n"
        "第三阶段（2026年9月1日至9月30日）：完成世界工具、语音条件能力、运行控制和错误回退。\n"
        "第四阶段（2026年10月1日至10月31日）：完成视觉 L1 证据链、规则式 L2 场景汇总、工作台展示和隐私边界。\n"
        "第五阶段（2026年11月1日至11月15日）：完成全量测试、项目表面审计、成果书整理和本机/局域网演示验收。"
    )


def deliverables() -> str:
    return (
        "1.作品表现形式\n"
        "（1）HutaoChatCore/HeadCore 源代码及必要配置示例；\n"
        "（2）Web Desk、原生微信小程序和本地视觉工作台；\n"
        "（3）自动化测试、运行手册、项目审计报告、项目计划书和毕业设计成果书。\n"
        "2.交付边界\n"
        "系统仅绑定 127.0.0.1:8000 进行本机或局域网演示，不提交密钥、模型权重、训练数据、运行日志和本地数据库。"
    )


def generate_task_document(output: Path) -> None:
    document = Document(str(TASK_TEMPLATE)) if TASK_TEMPLATE.is_file() else Document()
    configure_document(document)
    if not document.tables:
        table = document.add_table(rows=11, cols=14)
        table.style = "Table Grid"
    table = document.tables[0]
    set_cell_text(table.cell(0, 1), STUDENT)
    set_cell_text(table.cell(0, 5), STUDENT_ID)
    set_cell_text(table.cell(0, 9), MAJOR)
    set_cell_text(table.cell(0, 12), CLASS)
    set_cell_text(table.cell(1, 4), TITLE)
    set_cell_text(table.cell(2, 3), "软件开发类")
    set_cell_text(table.cell(2, 8), "社会实际")
    set_cell_text(table.cell(3, 3), "2026年6月30日至2026年11月15日")
    set_cell_text(table.cell(3, 13), SUPERVISOR)
    set_cell_text(table.cell(4, 2), "完成本机/局域网运行的 HeadCore 多模态角色陪伴系统，详见后附补充说明。")
    set_cell_text(table.cell(5, 2), "完成核心对话、客户端、认证、音频、受控世界工具和视觉 L1/L2，详见后附任务拆分。")
    set_cell_text(table.cell(6, 2), "2026年6月至11月分阶段完成需求、开发、视觉接线、测试和材料整理。")
    set_cell_text(table.cell(7, 2), "交付源代码、Web Desk、小程序、测试、运行手册、审计报告和毕业设计材料。")
    set_cell_text(table.cell(8, 2), "FastAPI、Pydantic、OpenCV、MediaPipe、Ultralytics、OWASP 等，详见后附参考资料。")
    set_cell_text(table.cell(9, 2), "\n\n教研室主任：                 年    月    日")
    set_cell_text(table.cell(10, 2), "\n\n院部负责人：                 年    月    日")
    add_task_supplement(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def result_sections() -> list[tuple[int, str, list[str]]]:
    return [
        (1, "选题说明", [
            "1.1 选题背景",
            "互联网陪伴与生成式人工智能的发展，使角色对话从单轮问答转向长期关系与多模态交互。但仅依赖模型提示词会产生人格漂移、记忆越权、外部事实混入和异常不可解释等问题。本课题选择 HeadCore 作为唯一认知主体，把客户端、模型、数据库、语音、视觉和世界工具都限制为能力提供者。",
            "1.2 市场与应用分析",
            "本系统面向个人本机或局域网演示，不定位为公网运营产品。它的应用价值是验证角色陪伴后端在身份一致性、证据边界、语音/视觉条件能力和安全控制方面的工程实现；相比普通聊天 Demo，重点放在可审计的运行时边界，而不是堆叠尚未验收的功能。",
            "1.3 项目目标",
            "完成文本聊天主链路、Web Desk、小程序客户端、账号与会话、HeadCore 人格和关系、记忆与世界证据、文件音频链路、视觉 L1 标签链与规则式 L2 场景汇总；所有已完成能力通过自动化测试或本机联调记录，未满足真实设备/外部服务条件的能力明确标为条件可用。",
        ]),
        (2, "需求分析", [
            "2.1 功能需求",
            "系统需要支持文本聊天、流式回复、OpenAI 兼容文本接口、受保护的 Web Desk、原生微信小程序、登录/注册/个人中心、文件语音识别、条件网页 TTS、记忆管理、天气/路线等受控世界工具、本地视觉工作台和运行状态查询。",
            "其中默认主线是文本聊天与 HeadCore；账号、数据库、SMTP、语音、世界工具和视觉均通过配置开关及运行条件控制。视觉工作台只允许管理员创建短时同意会话，原始帧不保存、不上传，人脸身份识别默认关闭。",
            "2.2 非功能需求",
            "安全性：认证使用会话、CSRF、密码哈希、限流和审计；敏感配置只从环境变量读取。可靠性：Provider 路由提供超时、重试、熔断和失败原因；上传接口限制大小、扩展名和临时文件生命周期。可维护性：模块按 HeadCore、Provider、存储、客户端和控制面分层。可验证性：每次改动配套测试与编译检查。",
            "2.3 明确排除项",
            "本成果不包含公网域名/HTTPS/反向代理、公开注册运营、实时视频问答、通用视觉问答、情绪/身份/意图视觉推断、历史平台适配器、旧版视觉方案和完整学习型世界模型。上述内容属于未纳入本次交付的范围。",
        ]),
        (3, "系统设计", [
            "3.1 总体架构",
            "客户端输入经过 FastAPI 统一入口，进入感知规范化、HeadCore 状态构建、决策与表达规划，再调用文本/语音/视觉/世界 Provider，最终通过对应客户端输出。HeadCore 是唯一认知主体，Provider 无权直接写入长期记忆或改变人格。",
            "图3-1 系统总体架构（文字化表示）",
            "客户端（Web Desk / 微信小程序 / OpenAI 兼容接口）\n    -> FastAPI Core（认证、聊天、音频、工作台、控制中心）\n    -> HeadCore（Self、关系、会话、记忆、决策、回复门禁）\n    -> Provider（DeepSeek / ASR / GPT-SoVITS / 世界工具 / 本地视觉）\n    -> 存储（JSONL 默认；MySQL、PostgreSQL、Qdrant 条件启用）",
            "3.2 模块设计",
            "HeadCore 模块负责认知状态与最终决策；persona/mind 负责胡桃人格、语气、关系与社会状态；services/providers 负责模型调用、质量评估和故障回退；audio/voice_chat 负责文件 ASR、情绪线索和网页 TTS；camera/workbench 负责明确同意的本地视觉会话、时序确认、L2 场景事实和证据展示；operations/control 负责运行状态、服务守护和验收。",
            "3.3 数据设计",
            "默认 JSONL 存储用于零外部依赖运行；条件启用时，MySQL Database V2 保存账号、档案、关系和控制审计，PostgreSQL 保存 Web 会话、消息和记忆，Qdrant 保存可选语义索引。视觉只保存结构化短时证据，不保存原始画面。",
            "3.4 详细设计",
            "视觉处理流程为：摄像头帧 -> OpenCV 采集 -> YOLO/MediaPipe 可选分析 -> 白名单标签 -> 时序确认 -> L2 场景事实汇总 -> 视觉问题证据块。规则式 L2 只输出 desk_work、desk_setup、street_vehicle、person_present 等有限状态，不推断情绪、身份和意图。",
            "3.5 安全与编码规范",
            "接口请求模型使用 Pydantic extra=forbid 和字段范围；危险能力采用全局开关、同意会话、来源许可和 CSRF 等多重门控；错误响应返回有限 reason_code，日志执行密钥脱敏；测试使用固定 fake provider 和临时目录隔离真实服务。",
        ]),
        (4, "系统实现", [
            "4.1 运行环境",
            "开发环境为 Windows 本机，Python 3.11，项目指定解释器为 D:\\Tool\\Progrmming-Tool\\anaconda\\envs\\new\\python.exe。核心服务绑定 127.0.0.1:8000；外部模型、数据库和语音服务均不随代码仓库提交。",
            "4.2 技术选型",
            "后端采用 FastAPI + Uvicorn + Pydantic；文本 Provider 采用 OpenAI 兼容协议；默认人格为 hutao_v1；认证使用 Argon2、Cookie 会话和 CSRF；数据库按身份与聊天职责分层；语音采用 FunASR/SenseVoice 与本地 GPT-SoVITS；视觉使用 OpenCV、可选 MediaPipe/Ultralytics，并通过白名单和时序确认控制证据。",
            "4.3 核心功能实现",
            "文本对话由 ChatService 组织上下文、HeadCore 决策、模型调用和回复评估；普通闲聊不再强制固定字数，只有用户明确要求短答时才启用长度门禁。文件音频接口有 25 MiB 默认限制、MIME/扩展名校验和 finally 清理。OpenAI 兼容接口明确拒绝当前未支持的图像/音频内容，避免静默丢失输入。",
            "工作台显示后端采集依赖、识别模型、诊断原因、浏览器预览和后端失败状态；控制中心的 camera_vision 状态区分未配置、降级和在线。当前本机默认关闭视觉，因此界面会诚实显示不可用，而不是伪造识别结果。",
        ]),
        (5, "系统测试", [
            "5.1 测试环境",
            "测试使用项目指定 Python 环境，pytest 运行 tests 目录并关闭 pytest 缓存；JavaScript 使用 Node.js 内置测试运行小程序客户端测试；前端工作台执行 node --check；Python 执行 compileall。",
            "5.2 测试结果",
            "截至本成果书生成日，后端全量测试为 929 passed、2 skipped；小程序测试为 5 passed；Python compileall、工作台 JavaScript 语法检查和 git diff --check 均通过。2 个跳过项属于未提供独立 MySQL 测试数据库的 opt-in 集成测试，不能解读为功能通过。",
            "5.3 代表性用例与缺陷处理",
        ]),
        (6, "实现成果", [
            "6.1 项目成果",
            "源码主线包括 app/main.py FastAPI 入口、app/head HeadCore、app/services ChatService、app/audio 文件音频、app/voice_chat TTS、app/camera 视觉 L1/L2、app/workbench 视觉工作台、app/operations 运行状态、frontend/site 落地页和 miniprogram 小程序。",
            "实现结果为：文本聊天主链路已验证；普通对话表达限制已自然化；音频上传边界和认证已补强；视觉工作台已有浏览器预览、后端能力诊断、采集失败提示和规则式场景事实。条件能力在缺配置时均 fail-closed。",
            "6.2 使用手册",
            "启动：在项目根目录执行 python -m uvicorn app.main:app --host 127.0.0.1 --port 8000；普通对话入口为 /desk，控制中心为 /control，视觉工作台为 /workbench。视觉工作台默认关闭，必须在本机配置管理员口令、视觉开关和真实模型/设备后再进行人工验收。",
            "6.3 交付限制",
            "本次交付不包含公网部署，不提供域名、HTTPS、备案、公开注册和远程摄像头上传。模型权重、训练素材、真实账号密钥、日志和运行数据库均不放入成果文档或代码交付包。",
        ]),
        (7, "总结与展望", [
            "7.1 工作总结",
            "本项目完成了从客户端、API、HeadCore、Provider、存储到控制面的可运行闭环。最主要的工程成果不是某一个模型，而是把人格、关系、记忆、外部证据、表达规划和安全门禁放入可测试的代码边界，并对条件能力保持诚实状态。",
            "7.2 问题与不足",
            "真实摄像头与视觉模型尚未在本机硬件上完成最终验收；MySQL/PostgreSQL/Qdrant、SMTP、DeepSeek、ASR/TTS 的完整生产级联调仍依赖外部服务；视觉 L2 目前是规则式状态汇总，不是通用图像问答；实时流式语音、共享限流、验证码、备份恢复和公网安全仍未完成。",
            "7.3 后续展望",
            "下一阶段优先完成本机摄像头与 YOLO/MediaPipe 固定版本验收、视觉状态评测集和失败重试；再评估本地小 VLM。公网相关安全工作必须在独立评审后进行，不纳入当前本机毕业设计主线。",
        ]),
    ]


REFERENCES = [
    "[1] Liang P, et al. Holistic Evaluation of Language Models (HELM)[EB/OL]. arXiv:2211.09110, 2022.",
    "[2] Park J S, et al. Generative Agents: Interactive Simulacra of Human Behavior[C]//UIST. 2023.",
    "[3] Packer C, et al. MemGPT: Towards LLMs as Operating Systems[EB/OL]. arXiv:2310.08560, 2023.",
    "[4] FastAPI. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/, 访问日期：2026-08-23.",
    "[5] Pydantic. Pydantic Documentation[EB/OL]. https://docs.pydantic.dev/, 访问日期：2026-08-23.",
    "[6] Uvicorn. Uvicorn Documentation[EB/OL]. https://www.uvicorn.org/, 访问日期：2026-08-23.",
    "[7] Python Software Foundation. Python 3.11 Documentation[EB/OL]. https://docs.python.org/3.11/, 访问日期：2026-08-23.",
    "[8] OpenCV. OpenCV Documentation[EB/OL]. https://docs.opencv.org/, 访问日期：2026-08-23.",
    "[9] Google. MediaPipe Solutions Documentation[EB/OL]. https://ai.google.dev/edge/mediapipe/solutions, 访问日期：2026-08-23.",
    "[10] Ultralytics. Ultralytics YOLO Documentation[EB/OL]. https://docs.ultralytics.com/, 访问日期：2026-08-23.",
    "[11] OWASP Foundation. OWASP Application Security Verification Standard 4.0.3[S/OL]. 2023.",
    "[12] HutaoChatCore. HUTAOCHATCORE_COMPLETE_ARCHITECTURE_AND_ACCEPTANCE_MANUAL[EB/OL]. 项目文档，2026.",
]


def add_result_table(document: Document, level: int) -> None:
    if level == 2:
        add_table(
            document,
            ["需求类别", "当前实现", "状态", "边界"],
            [
                ["核心对话", "文本聊天、流式回复、HeadCore 状态和回复评估", "已完成", "文本 Provider 需配置"],
                ["账号与会话", "注册、登录、邮箱验证、密码重置、CSRF、限流", "已完成主链路", "邮件服务条件可用"],
                ["音频能力", "文件上传、ASR、情绪线索、网页 TTS", "条件可用", "依赖本地模型和服务"],
                ["视觉能力", "L1 标签证据链、L2 规则式场景状态、工作台", "L1/L2 已完成", "真实设备仍需验收"],
                ["运行与安全", "控制中心、状态诊断、fail-closed 和审计", "已完成", "只面向本机/局域网"],
            ],
        )
    elif level == 3:
        add_table(
            document,
            ["层次", "模块/路径", "职责", "数据边界"],
            [
                ["客户端", "Web Desk / 小程序 / OpenAI 兼容接口", "输入、会话和结果展示", "不决定人格和权限"],
                ["核心服务", "app/main.py / ChatService", "认证、上下文、决策和输出", "统一入口、失败可解释"],
                ["HeadCore", "app/head、persona、mind", "Self、关系、记忆和回复门禁", "Provider 无权越权写长期记忆"],
                ["能力提供者", "providers、audio、camera、world", "模型、语音、视觉和外部证据", "默认关闭或条件启用"],
                ["存储与控制", "storage、operations、workbench", "JSONL/条件数据库、状态和审计", "密钥和原始帧不入文档"],
            ],
        )
    elif level == 4:
        add_table(
            document,
            ["实现模块", "关键文件", "实现结果", "当前状态"],
            [
                ["对话主链路", "app/services/chat_service.py", "上下文组织、HeadCore 决策、Provider 调用和回复评估", "已完成"],
                ["文件音频", "app/audio/file_service.py", "25 MiB 限制、MIME/扩展名校验、临时文件清理", "已完成"],
                ["视觉证据", "app/camera/evidence_store.py", "时序确认标签、TTL、场景状态和对话证据块", "已完成"],
                ["视觉工作台", "app/workbench/router.py、static/workbench", "能力诊断、预览、场景状态和失败提示", "条件可用"],
                ["运行控制", "app/operations、app/control", "健康状态、服务守护和本地验收入口", "已完成主线"],
            ],
        )
    elif level == 6:
        add_table(
            document,
            ["交付物", "路径", "内容", "是否纳入提交"],
            [
                ["后端源代码", "app/", "FastAPI、HeadCore、Provider、存储和控制面", "是"],
                ["客户端", "frontend/、miniprogram/", "Web Desk、PWA 和原生小程序", "是"],
                ["测试与脚本", "tests/、scripts/", "自动化测试、审计和运行工具", "是"],
                ["项目文档", "docs/", "技术计划、复盘、审计和毕业设计材料", "是"],
                ["本地资产", "data/models/、external/", "模型权重、训练素材和运行数据", "否"],
            ],
        )
    elif level == 7:
        add_table(
            document,
            ["问题/限制", "当前事实", "后续安排"],
            [
                ["视觉设备验收", "真实摄像头和固定模型尚未完成最终硬件验收", "本机条件具备后执行人工验收"],
                ["外部服务依赖", "数据库、SMTP、DeepSeek、ASR/TTS 需按配置启用", "继续保留 fail-closed 和联调记录"],
                ["视觉理解范围", "L2 是规则式场景汇总，不是通用图像问答", "评估本地小 VLM，单独立项"],
                ["公网安全", "未做域名、HTTPS、公开注册和生产反向代理", "公网前单独完成 P0 安全评审"],
            ],
        )


def add_result_abstract(document: Document) -> None:
    document.add_heading("摘要", level=1)
    add_body(
        document,
        "本项目实现了一个以 HeadCore 为唯一认知主体的本地多模态角色陪伴系统。系统通过 FastAPI 统一接入 Web Desk、原生微信小程序和 OpenAI 兼容文本接口，在文本聊天主链路上组织人格、关系、记忆、外部证据和回复门禁；文件音频、语音合成、数据库、世界工具和视觉能力作为受控 Provider 按配置启用。视觉部分完成白名单标签证据链和规则式 L2 场景事实汇总，并在工作台展示能力状态、诊断原因和置信度。项目当前仅用于本机或局域网演示，不部署公网。",
    )
    add_body(document, "关键词：HeadCore；FastAPI；角色陪伴；证据边界；视觉工作台；规则式场景状态")


def generate_result_document(output: Path) -> None:
    document = Document(str(RESULT_TEMPLATE)) if RESULT_TEMPLATE.is_file() else Document()
    clear_body(document)
    configure_document(document)
    add_title(document, TITLE)
    add_center(document, "毕业设计成果书", size=16, bold=True)
    add_center(document, "软件学院毕业设计成果材料", size=12)
    add_metadata_table(document)
    add_center(document, DATE, size=11)
    document.add_page_break()

    add_result_abstract(document)
    document.add_page_break()

    add_center(document, "目录", size=16, bold=True)
    add_body(document, "目录按学校成果规范组织；在 Word 中打开后可使用“更新域”刷新页码。")
    add_table(
        document,
        ["章节", "内容"],
        [
            ["1 选题说明", "背景、应用价值和项目目标"],
            ["2 需求分析", "功能、非功能需求与交付边界"],
            ["3 系统设计", "总体架构、模块、数据和安全设计"],
            ["4 系统实现", "运行环境、技术选型和关键实现"],
            ["5 系统测试", "测试环境、结果和代表性用例"],
            ["6 实现成果", "源码、客户端、文档和交付范围"],
            ["7 总结与展望", "完成情况、限制和后续工作"],
            ["参考文献 / 附录", "参考资料、运行命令和交付边界"],
        ],
    )
    document.add_page_break()

    for index, (level, heading, content) in enumerate(result_sections()):
        if index:
            document.add_page_break()
        document.add_heading(f"{level} {heading}", level=1)
        for item in content:
            if item.startswith(f"{level}."):
                document.add_heading(item, level=2)
            elif item.startswith("图3-1"):
                document.add_heading(item, level=2)
            elif "->" in item or "客户端（" in item:
                add_code(document, item)
            else:
                add_body(document, item)
        add_result_table(document, level)
        if heading == "系统测试":
            add_table(
                document,
                ["测试层次", "命令/范围", "结果", "真实性边界"],
                [
                    ["后端回归", "pytest tests", "929 passed, 2 skipped", "契约和确定性逻辑；2 项 MySQL opt-in 未执行"],
                    ["语法检查", "compileall app scripts tests", "PASS", "不等同于真实外部服务验收"],
                    ["小程序", "node --test 两个测试文件", "5 passed", "客户端状态与 API 边界"],
                    ["视觉", "规则状态/工作台测试", "PASS", "真实摄像头和模型仍需人工验收"],
                ],
            )
    document.add_page_break()
    document.add_heading("参考文献", level=1)
    for item in REFERENCES:
        add_bullet(document, item)
    document.add_page_break()
    document.add_heading("附录", level=1)
    document.add_heading("附录A 运行命令", level=2)
    add_code(document, "D:\\Tool\\Progrmming-Tool\\anaconda\\envs\\new\\python.exe -m uvicorn app.main:app --host 127.0.0.1 --port 8000")
    add_code(document, "D:\\Tool\\Progrmming-Tool\\anaconda\\envs\\new\\python.exe -m pytest tests -q -p no:cacheprovider")
    document.add_heading("附录B 交付文件边界", level=2)
    add_body(document, "提交代码时保留 app、scripts、tests、frontend、miniprogram、migrations、deploy 和必要文档；排除 .env、模型权重、训练数据、运行日志、临时文件、生成音频和本地数据库。")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def markdown_task() -> str:
    description = task_description().replace("\n", "\n\n")
    requirements = task_requirements().replace("\n", "\n\n")
    schedule = implementation_schedule().replace("\n", "\n\n")
    output = deliverables().replace("\n", "\n\n")
    return f"""# {TITLE}：毕业设计任务书/项目计划书

> 生成日期：{DATE}。本文件依据学校《2027届学生毕业设计任务书》模板整理，个人信息待填写。

| 字段 | 内容 |
| --- | --- |
| 学生姓名 | {STUDENT} |
| 学号 | {STUDENT_ID} |
| 专业 | {MAJOR} |
| 班级 | {CLASS} |
| 课题类型 | 软件开发类 |
| 课题来源 | 社会实际 |
| 指导教师 | {SUPERVISOR} |
| 时间 | 2026年6月30日至2026年11月15日 |

## 课题说明

{description}

## 真实性边界

- 当前只面向本机/局域网演示，未部署公网。
- 当前已实现文本聊天、HeadCore、Web Desk、小程序、认证边界、文件音频、条件 TTS、视觉 L1 和规则式 L2 场景事实汇总。
- 视觉真实摄像头、YOLO/MediaPipe 固定模型、外部数据库与生产级邮件/模型服务仍需条件验收。
- 不纳入本课题成果：历史平台适配器、未验收的通用视觉问答、公开运营和未完成的公网安全设施。

## 任务拆分

{requirements}

## 实施步骤与时间安排

{schedule}

## 交付物

{output}
"""


def markdown_result() -> str:
    lines = [f"# {TITLE}：毕业设计成果书", "", f"> 生成日期：{DATE}。本成果书依据软件学院毕业设计成果参考规范整理。", "", "## 摘要", "", "本项目实现了一个以 HeadCore 为唯一认知主体的本地多模态角色陪伴系统，完成文本聊天、客户端、认证边界、条件音频能力、视觉 L1/L2 证据链和工作台状态展示。系统仅用于本机或局域网演示，不部署公网。", "", "关键词：HeadCore；FastAPI；角色陪伴；证据边界；视觉工作台；规则式场景状态", ""]
    for level, heading, content in result_sections():
        lines.extend([f"## {level} {heading}", ""])
        for item in content:
            if item.startswith(f"{level}."):
                lines.extend([f"### {item}", ""])
            elif "->" in item or item.startswith("客户端（"):
                lines.extend(["```text", item, "```", ""])
            else:
                lines.extend([item, ""])
    lines.extend(["## 参考文献", ""])
    lines.extend(f"{item}  " for item in REFERENCES)
    lines.extend(["", "## 附录", "", "- Python 项目测试：929 passed、2 skipped。", "- 小程序 Node 测试：5 passed。", "- 公网部署、真实视觉设备和外部生产服务不在当前交付范围。", ""])
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate truthful graduation-design documents from the project state.")
    parser.add_argument("--output-dir", type=Path, default=OUTPUT_DIR)
    args = parser.parse_args()
    output_dir = args.output_dir if args.output_dir.is_absolute() else PROJECT_ROOT / args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    suffix = "排版优化版_2026-08-23"
    generate_task_document(output_dir / f"HeadCore_毕业设计任务书_项目计划书_{suffix}.docx")
    generate_result_document(output_dir / f"HeadCore_毕业设计成果书_{suffix}.docx")
    (output_dir / f"HeadCore_毕业设计任务书_项目计划书_{suffix}.md").write_text(markdown_task(), encoding="utf-8")
    (output_dir / f"HeadCore_毕业设计成果书_{suffix}.md").write_text(markdown_result(), encoding="utf-8")
    print(output_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
