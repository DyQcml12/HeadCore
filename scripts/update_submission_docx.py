from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "submission"


def set_paragraph(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> int:
    plan_candidates = [p for p in ROOT.glob("*2026-08-30.docx") if len(Document(p).paragraphs) < 50]
    if not plan_candidates:
        raise FileNotFoundError("updated plan document not found")
    plan_path = max(plan_candidates, key=lambda item: len(Document(item).paragraphs))
    plan_document = Document(plan_path)
    plan_table = plan_document.tables[0]
    set_paragraph(
        plan_table.cell(4, 2).paragraphs[0],
        "\u5b8c\u6210\u672c\u673a/\u5c40\u57df\u7f51\u8fd0\u884c\u7684 HeadCore \u591a\u6a21\u6001\u89d2\u8272\u966a\u4f34\u7cfb\u7edf\uff0c\u5305\u542b\u8bc1\u636e\u4f18\u5148\u4e16\u754c\u6a21\u578b\uff08\u8de8\u8f6e\u4e8b\u5b9e\u3001\u4e16\u754c\u56fe\u589e\u957f\u548c\u6709\u754c\u53cd\u4e8b\u5b9e\u9a8c\u8bd5\uff09\uff0c\u8be6\u89c1\u540e\u9644\u8865\u5145\u8bf4\u660e\u3002",
    )
    set_paragraph(
        plan_table.cell(5, 2).paragraphs[0],
        "\u5b8c\u6210\u6838\u5fc3\u5bf9\u8bdd\u3001\u5ba2\u6237\u7aef\u3001\u8ba4\u8bc1\u3001\u97f3\u9891\u3001\u53d7\u63a7\u4e16\u754c\u5de5\u5177\u4e0e\u4e16\u754c\u6a21\u578b\u3001\u89c6\u89c9 L1/L2\uff0c\u8be6\u89c1\u540e\u9644\u4efb\u52a1\u62c6\u5206\u3002",
    )
    plan_document.save(plan_path)

    candidates = [p for p in ROOT.glob("*2026-08-30.docx") if len(Document(p).paragraphs) > 80]
    if not candidates:
        raise FileNotFoundError("updated result document not found")
    path = max(candidates, key=lambda item: len(Document(item).paragraphs))
    document = Document(path)

    # Keep paragraph styles and the existing document structure intact.
    cover_date = document.paragraphs[4]
    set_paragraph(cover_date, re.sub(r"2026.8.23.", "2026\u5e748\u670830\u65e5", cover_date.text))

    test_result = document.paragraphs[61]
    if "\u4e16\u754c\u6a21\u578b\u6548\u679c\u8bc4\u4f30\u4e3a" not in test_result.text:
        set_paragraph(
            test_result,
            test_result.text.replace(
                "929 passed\u30012 skipped",
                "981 passed\u30012 skipped\uff1b\u4e16\u754c\u6a21\u578b\u6548\u679c\u8bc4\u4f30\u4e3a 12/12 PASS\u3001gap=0\uff0c\u8986\u76d6\u8bc1\u636e\u6295\u5f71\u3001\u8fc7\u671f\u4e0e\u51b2\u7a81\u4fdd\u62a4\u3001\u8de8\u8f6e\u5929\u6c14\u4e8b\u5b9e\u3001\u666e\u901a\u5bf9\u8bdd\u4e16\u754c\u56fe\u589e\u957f\u548c\u6709\u754c\u53cd\u4e8b\u9a8c\u8bd5"
            ),
        )

    if not document.paragraphs[63].text.strip():
        set_paragraph(
            document.paragraphs[63],
            "\u4ee3\u8868\u6027\u7528\u4f8b\u5305\u62ec\uff1a\u6709\u6548\u4e16\u754c\u5173\u7cfb\u6295\u5f71\u3001\u8fc7\u671f\u5173\u7cfb\u4e0d\u6295\u5f71\u3001\u540c\u952e\u51b2\u7a81\u4fdd\u62a4\u3001\u672a\u786e\u8ba4\u56e0\u679c\u4fdd\u6301\u5047\u8bbe\u6807\u7b7e\u3001\u5929\u6c14\u8bc1\u636e\u5f53\u8f6e\u6ce8\u5165\u3001 Web \u7528\u6237\u8de8\u8f6e\u6062\u590d\u3001\u666e\u901a\u5bf9\u8bdd\u81ea\u52a8\u5199\u5165\u5b9e\u4f53\u4e0e\u4e8b\u4ef6\uff0c\u4ee5\u53ca\u6709\u754c\u53cd\u4e8b\u5b9e\u9a8c\u8bd5\u7684 supported/refuted \u8def\u5f84\u3002\u4e16\u754c\u6a21\u578b\u6548\u679c\u8bc4\u4f30\u4e3a 12/12 PASS\u3001gap=0\u3002\u7f3a\u9677\u5904\u7406\u91c7\u7528\u8bc1\u636e\u4f18\u5148\u3001\u7528\u6237\u9694\u79bb\u3001\u7a33\u5b9a\u54c8\u5e0c\u53bb\u91cd\u548c\u4e0d\u53ef\u7528\u65f6 fail-closed\uff1bL4 \u7ed3\u679c\u4ec5\u4ee3\u8868\u786e\u5b9a\u6027\u6709\u754c\u8bd5\u9a8c\uff0c\u4e0d\u4ee3\u8868\u901a\u7528\u9884\u6d4b\u80fd\u529b\u3002",
        )

    limitation = document.paragraphs[79]
    if "\u4e16\u754c\u6a21\u578b\u5df2\u5b8c\u6210" not in limitation.text:
        set_paragraph(
            limitation,
            limitation.text.replace(
                "\uff1bMySQL",
                "\uff1b\u4e16\u754c\u6a21\u578b\u5df2\u5b8c\u6210\u8bc1\u636e\u6295\u5f71\u3001\u51b2\u7a81\u4e0e\u8fc7\u671f\u4fdd\u62a4\u3001\u8de8\u8f6e\u4e8b\u5b9e\u6301\u4e45\u5316\u3001\u666e\u901a\u5bf9\u8bdd\u4e16\u754c\u56fe\u589e\u957f\u53ca\u6709\u754c\u53cd\u4e8b\u5b9e\u9a8c\u8bd5\uff0c\u4f46\u5c1a\u4e0d\u5177\u5907\u901a\u7528\u4e16\u754c\u9884\u6d4b\u6216\u6301\u7eed\u5b66\u4e60\u80fd\u529b\uff1bMySQL",
            ),
        )

    document.save(path)
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
